import streamlit as st
import json
import io
from docx import Document
from groq import Groq
from openai import OpenAI
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import Chroma

# ==============================================================================
# 1. PROVIDER & LOCAL DATABASE CONFIGURATIONS
# ==============================================================================
# Try to get key from Streamlit secrets (Cloud), fallback to local if it fails
try:
    GROQ_API_KEY = st.secrets["GROQ_API_KEY"]
except:
    GROQ_API_KEY = "YOUR_LOCAL_GROQ_API_KEY_HERE" # Put your key here for local testing

groq_client = Groq(api_key=GROQ_API_KEY)
local_client = OpenAI(base_url="http://localhost:11434/v1", api_key="ollama")

DB_FOLDER = "C:/Users/ngunj/Desktop/MwalimuApp/chroma_db"

@st.cache_resource
def load_vector_db():
    embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
    return Chroma(persist_directory=DB_FOLDER, embedding_function=embeddings)

try:
    vector_db = load_vector_db()
except Exception as e:
    st.error(f"Could not load Chroma database. Did you run embed_syllabus.py? Error: {e}")

# ==============================================================================
# 2. WORD DOCUMENT GENERATOR HELPER
# ==============================================================================
def create_word_document(lesson_data):
    doc = Document()
    doc.add_heading(lesson_data.get('lesson_title', 'KICD Lesson Plan'), 0)
    
    doc.add_heading('Basic Information', level=1)
    doc.add_paragraph(f"Subject: {lesson_data.get('subject', 'N/A')}")
    doc.add_paragraph(f"Class Level: {lesson_data.get('class_level', 'N/A')}")
    doc.add_paragraph(f"Duration: {lesson_data.get('duration', 'N/A')}")
    doc.add_paragraph(f"Strand: {lesson_data.get('strand', 'N/A')} ({lesson_data.get('sub_strand', 'N/A')})")
    
    doc.add_heading('🎯 Learning Outcomes', level=1)
    for outcome in lesson_data.get("learning_outcomes", []):
        doc.add_paragraph(outcome, style='List Bullet')
        
    doc.add_heading('🛠️ Learning Resources', level=1)
    resources = lesson_data.get("learning_resources", [])
    if isinstance(resources, list):
        doc.add_paragraph(", ".join(resources))
    else:
        doc.add_paragraph(str(resources))
        
    doc.add_heading('🛑 Classroom Management Tips', level=1)
    for tip in lesson_data.get("classroom_management_tips", []):
        doc.add_paragraph(tip, style='List Bullet')
        
    doc.add_heading('🚶‍♂️ Lesson Procedures', level=1)
    for proc in lesson_data.get("lesson_procedures", []):
        doc.add_heading(f"{proc.get('stage')} ({proc.get('time')})", level=2)
        doc.add_paragraph(f"Teacher Activity: {proc.get('teacher_activity')}")
        doc.add_paragraph(f"Learner Activity: {proc.get('learner_activity')}")
        
    doc.add_heading('📝 Assessment', level=1)
    doc.add_paragraph(str(lesson_data.get('assessment', 'N/A')))
    
    doc.add_heading('💭 Reflection', level=1)
    doc.add_paragraph(str(lesson_data.get('reflection', 'N/A')))
    
    # Save document to a bytes buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ==============================================================================
# 3. DYNAMIC SYSTEM INSTRUCTIONS 
# ==============================================================================
SYSTEM_INSTRUCTIONS = """
You are MwalimuAI, a Kenya CBC lesson plan generator. You MUST follow KICD format 100%.
You will be provided with official curriculum snippets extracted via RAG in the user prompt. Use them to maintain absolute compliance.

STRICT RULES:
1. Output ONLY valid JSON. No explanations.
2. Include ALL fields. If missing, use "TBD".
3. Learning outcomes MUST start with: "By the end of the sub strand, the learner should be able to: " + action verb.
4. Duration: Grade 1-3 = 30 min, Grade 4-6 = 35 min, Grade 7 = 40 min.
5. Core competencies: pick 2-3 EXACTLY from [Communication and Collaboration, Critical Thinking and Problem Solving, Creativity and Imagination, Digital Literacy, Learning to Learn, Self-Efficacy, Citizenship].
6. Values: pick 1-2 EXACTLY from [Respect, Responsibility, Unity and Peace, Love and Patriotism, Social Justice and Integrity].
7. PCIs: pick 1-2 EXACTLY from [Health and Hygiene, Life Skills and Self-Awareness, Citizenship and Social Cohesion, Environmental Awareness, Financial Literacy, Safety and Security].
8. Learning resources must be locally available.
9. Lesson procedures: Introduction (5-7min), Development (20-25min), Conclusion (8-10min). Teacher facilitates; Learner is active/experiential.
10. Assessment tests outcomes directly.
11. Reflection: 2 sentences on what worked/what to improve.
12. Provide 3 actionable "Classroom Management and Control Tips".
13. If subject is Kiswahili, generated text MUST be in Kiswahili.

JSON SCHEMA TO FOLLOW:
{
  "subject": "", "strand": "", "sub_strand": "", "class_level": "", "date": "", "duration": "", "lesson_title": "",
  "learning_outcomes": [], "core_competencies": [], "values": [], "pcis": [], "learning_resources": [],
  "lesson_procedures": [{"stage": "Introduction", "time": "", "teacher_activity": "", "learner_activity": ""}],
  "assessment": "", "reflection": "", "classroom_management_tips": []
}
"""

def clean_json_output(raw_text):
    text = raw_text.strip()
    if text.startswith("```json"):
        text = text[7:]
    if text.startswith("
```"):
        text = text[3:]
    if text.endswith("```"):
        text = text[:-3]
    return text.strip()

# ==============================================================================
# 4. STREAMLIT FRONTEND USER INTERFACE
# ==============================================================================
st.title(" 🇰🇪 MwalimuAI Prototype (True RAG Mode)")
st.subheader("Dynamic KICD Curriculum Context Retrieval")

connection_mode = st.radio("Select Network Mode:", ["🌐 Online Mode (Lightning Fast via Groq)", "🔌 Offline Mode (Local Qwen-7B via Ollama)"])
st.markdown("---")

class_level = st.selectbox("Select Class Level", ["Grade 1", "Grade 2", "Grade 3", "Grade 4", "Grade 5", "Grade 6", "Grade 7"])
subject = st.text_input("Enter Subject (e.g., Mathematics, English, Agriculture and Nutrition)")
topic = st.text_input("Enter Topic / Strand (e.g., Algebra, Linear Equations, Pre-reading)")

if st.button("Generate Lesson Plan"):
    if subject and topic:
        with st.spinner("🔍 Querying local KICD database and generating lesson plan..."):
            search_query = f"Class Level: {class_level}, Subject: {subject}, Topic/Strand: {topic}"
            docs = vector_db.similarity_search(search_query, k=4)
            retrieved_context = "\n---\n".join([doc.page_content for doc in docs])
            
            user_prompt = f"USER REQUEST:\nClass Level: {class_level}\nSubject: {subject}\nTopic/Strand: {topic}\n\nOFFICIAL RETRIEVED CURRICULUM CONTEXT FROM LOCAL PDFs:\n{retrieved_context}"
            
            try:
                raw_json = ""
                if "Online Mode" in connection_mode:
                    response = groq_client.chat.completions.create(
                        model="llama-3.3-70b-versatile",
                        messages=[
                            {"role": "system", "content": SYSTEM_INSTRUCTIONS},
                            {"role": "user", "content": user_prompt}
                        ],
                        response_format={"type": "json_object"},
                        temperature=0.0
                    )
                    raw_json = response.choices[0].message.content
                else:
                    response = local_client.chat.completions.create(
                        model="qwen:7b",
                        messages=[
                            {"role": "system", "content": SYSTEM_INSTRUCTIONS},
                            {"role": "user", "content": user_prompt}
                        ],
                        temperature=0.0
                    )
                    raw_json = clean_json_output(response.choices[0].message.content)
                
                lesson_data = json.loads(raw_json)
                
                st.success("Lesson Plan Generated via Dynamic PDF Context Lookup!")
                
                # --- EXPORT TO WORD BUTTON ---
                word_buffer = create_word_document(lesson_data)
                safe_filename = f"{lesson_data.get('subject', 'Subject').replace(' ', '_')}_{lesson_data.get('class_level', 'Level').replace(' ', '')}_LessonPlan.docx"
                
                st.download_button(
                    label="📄 Download Lesson Plan as Word Document",
                    data=word_buffer,
                    file_name=safe_filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    type="primary"
                )
                st.markdown("---")
                
                # Render the UI
                st.markdown(f"### 📋 {lesson_data.get('lesson_title', 'Lesson Plan')}")
                col1, col2 = st.columns(2)
                with col1:
                    st.write(f"**Subject:** {lesson_data.get('subject')}")
                    st.write(f"**Class Level:** {lesson_data.get('class_level')}")
                with col2:
                    st.write(f"**Duration:** {lesson_data.get('duration')}")
                    st.write(f"**Strand:** {lesson_data.get('strand')} ({lesson_data.get('sub_strand')})")
                
                st.write("#### 🎯 Learning Outcomes")
                for outcome in lesson_data.get("learning_outcomes", []): st.write(f"- {outcome}")
                    
                st.write("#### 🛠️ Learning Resources")
                resources = lesson_data.get("learning_resources")
                st.write(", ".join(resources) if isinstance(resources, list) else resources)
                
                st.write("#### 🛑 Classroom Management & Control Tips")
                for idx, tip in enumerate(lesson_data.get("classroom_management_tips", []), 1): st.info(f"**Tip {idx}:** {tip}")
                
                st.write("#### 🚶‍♂️ Lesson Procedures")
                for proc in lesson_data.get("lesson_procedures", []):
                    with st.expander(f"{proc.get('stage')} ({proc.get('time')})"):
                        st.write(f"**Teacher Activity:** {proc.get('teacher_activity')}")
                        st.write(f"**Learner Activity:** {proc.get('learner_activity')}")
                        
                st.write("#### 📝 Assessment")
                st.write(lesson_data.get("assessment"))
                
                st.write("#### 💭 Reflection")
                st.write(lesson_data.get("reflection"))
                
                with st.expander("👁️ View Raw PDF Context Pulled by RAG"):
                    st.text(retrieved_context)
                        
            except Exception as e:
                st.error(f"An error occurred. Check your terminal logs. Error details: {e}")
    else:
        st.warning("Please fill in both the Subject and Topic fields.")
